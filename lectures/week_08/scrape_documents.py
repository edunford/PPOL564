# !pip install PyPDF2
# !pip install python-docx

import PyPDF2 # Scraping PDFs
import docx # Scraping Word Documents


import warnings
warnings.filterwarnings("ignore")


# %% -----------------------------------------

#############################################
########## Scraping TEXT in a PDF ###########
#############################################

# Mercy Corp Report As example
# https://www.mercycorps.org/sites/default/files/2019-11/Motivations%20and%20Empty%20Promises_Mercy%20Corps_Full%20Report_0.pdf

pdfFileObj = open('mercy_corp.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
pdfReader.numPages
pdfReader.isEncrypted
pageObj = pdfReader.getPage(10)
print(pageObj.extractText())
pdfFileObj.close()


# %% -----------------------------------------

# Extract all the text content in the PDF

def read_pdf(file):
    with open(file, 'rb') as pdfFileObj:
        # Open the pdf
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj)

        # Locate the number of pages
        n_pages = pdfReader.numPages

        # Loop through the pages and store the content by
        # appending to a string
        content = ""
        for i in range(n_pages):
            content += pdfReader.getPage(i).extractText()
    return content

# Examine the content
mc_content = read_pdf("mercy_corp.pdf")
print(mc_content)


# %% -----------------------------------------

# CAUTION! Not all PDFs are equal. Some are really difficult to parse.
# Given that the spaces aren't special characters.
tw_content = read_pdf("thomas_wood.pdf")
print(tw_content)

# When you run into this, you'll have to think through an alternative parsing strategy. No free lunch here.


# %% -----------------------------------------

################################################
########## Scraping TEXT in WORD doc ###########
################################################

doc = docx.Document("Easterly_and_Levine.docx")

len(doc.paragraphs)

doc.paragraphs[0].text

for i in doc.paragraphs:
    print(i.text)


# %% -----------------------------------------

# Wrap into a function

def get_word(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(get_word("Easterly_and_Levine.docx"))
